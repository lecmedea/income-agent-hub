#!/usr/bin/env python3
"""Generate Income Agent strategy document (foundation for full playbook)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path.home() / "Downloads" / f"income-agent-strategy-{date.today().isoformat()}.docx"

PARTS = [
    ("Часть I. Видение: ребёнок-агент", [
        "Income Agent — не скрипт, а наследник знаний оператора и инструментов Grok/Codex.",
        "Цель: легальный, повторяемый денежный поток для ИП Бежаев С.В.",
        "Принцип заботы: передавать навыки, журналы, KPI — как семейное наследие системы.",
        "Граница: человек подписывает, человек несёт юридическую ответственность.",
    ]),
    ("Часть II. Архитектура 2026", [
        "L0: hunt-lite.py + n8n + Telegram @Agent00AI_bot (0 токенов).",
        "L1: propose откликов через Grok по запросу оператора.",
        "L2: deliver через site-builder, docx, pptx, russian-marketing-2026.",
        "L3: счёт/акт через ЭДО — черновик агентом, подпись оператором.",
        "Hub API :8765, SQLite, очередь Google Sheets, Desktop Income Agent 007.",
    ]),
    ("Часть III. Направления заработка", [
        "1. Сайты и лендинги (45–80k ₽) — азимут, grillz кейсы.",
        "2. AI-чат и n8n автоматизации для SMB.",
        "3. SMM + shorts 50% mix (тренды Чижова 2026).",
        "4. SEO-старт и CIS-индексация.",
        "5. AI-инфлюенсеры и контент-заводы.",
        "6. Part-time PM / AI-маркетолог (hh remote).",
        "7. Upsell текущих: KAIFOOI, Грантек, Grillz, Азимут.",
    ]),
    ("Часть IV. ЭДО и подписи", [
        "Агент готовит: КП, ТЗ, акт, счёт — шаблоны в docx skill.",
        "Оператор: подпись КЭП/ЭДО в личном кабинете (Сбер, Такском, и т.д.).",
        "В таблице Finance — колонка «ЭДО/счёт» для каждой сделки.",
        "Не автоматизировать подпись без HSM-токена оператора.",
    ]),
    ("Часть V. KPI и стадии заказа", [
        "Стадии: new → review → proposed → approved → in_progress → grok/codex → delivered → invoiced → paid.",
        "KPI недели: лиды≥20, отклики≥5, win-rate≥10%, средний чек≥минимумов прайса.",
        "Лист KPI в Google Sheets считает формулы автоматически.",
    ]),
    ("Часть VI. Skills roadmap", [
        "Уже: income-agent, schedule, autonomy, russian-marketing-2026.",
        "Далее: income-agent-kwork (API/парсер), income-agent-edo (шаблоны счетов),",
        "income-agent-portfolio-sync (Platforms лист), income-agent-codex-bridge.",
        "Codex: зеркальные skills в ~/.codex/skills для deliver.",
    ]),
    ("Часть VII. Маркетинг РФ 2026 (Чижов)", [
        "Дофаномика → 50% shorts, shoppertainment, смелая коммуникация.",
        "Цифровая депривация → честность, «настоящий» голос бренда.",
        "ИИ-прогресс → AI как усилитель, не замена эксперта.",
        "Поколения → разный тон для Z / миллениалов / B2B.",
        "Визуал → киношная эстетика для Bond-desktop, тёплый тон для клиник.",
    ]),
    ("Часть VIII. Риски и этика", [
        "Запрет: имитация оператора без согласия, серые схемы, спам.",
        "Обязательно: журнал лидов, approve перед отправкой, 152-ФЗ для персональных данных.",
        "Медицина/клиники: комплаенс, без диагнозов в AI-ответах.",
    ]),
    ("Часть IX. 90-дневный план", [
        "Дни 1–14: Hub + бот + Sheets + 10 откликов FL.ru.",
        "Дни 15–45: 2 закрытых проекта, публикация skills на GitHub.",
        "Дни 46–90: входящие с vc/TG, автодайджест, первый recurring клиент.",
    ]),
    ("Часть X. Будущее агента (2027+)", [
        "Полуавтономный propose с human-in-the-loop 30 секунд.",
        "Интеграция Kwork/hh API при появлении ключей.",
        "Единый дашборд P&L ИП в Sheets + выгрузка в бухгалтерию.",
        "«Ребёнок» учится: каждый выигранный тендер → новый шаблон в skills.",
    ]),
]

# Expand each part with detailed subsections for volume
SUB_TOPICS = [
    "Экономика внимания и конверсия лидов",
    "Скрипты откликов A/B",
    "Пакетирование услуг",
    "Работа с возражениями",
    "Шаблоны смет",
    "Контроль дедлайнов",
    "Ретеншн клиентов",
    "Реферальная программа",
    "Налоговый учёт ИП (обзор)",
    "Резервный фонд 3 месяца",
]


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)

    title = doc.add_heading("СТРАТЕГИЯ INCOME AGENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Полевой документ развития автономного агента заработка")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Оператор: ИП Бежаев Сергей Викторович · {date.today().strftime('%d.%m.%Y')}")
    doc.add_page_break()

    doc.add_heading("Оглавление (авто-структура)", level=1)
    for i, (part_title, _) in enumerate(PARTS, 1):
        doc.add_paragraph(f"{i}. {part_title}", style="List Number")
    doc.add_page_break()

    for part_title, bullets in PARTS:
        doc.add_heading(part_title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
        for st in SUB_TOPICS:
            doc.add_heading(st, level=2)
            doc.add_paragraph(
                f"В контексте «{part_title}»: {st}. "
                "Агент фиксирует гипотезу в Sheets, тестирует на 1–2 лидах, "
                "измеряет KPI (отклик→созвон→сделка), масштабирует только после approve оператора. "
                "Связка с Grok skills: income-agent, russian-marketing-2026, income-agent-autonomy."
            )
            doc.add_paragraph(
                "Практика: еженедельный ретро в Telegram /desktop, обновление Platforms листа, "
                "корректировка минимальных ставок из operator-profile.md при инфляции спроса."
            )
        doc.add_page_break()

    doc.add_heading("Приложение A. Реквизиты и ссылки", level=1)
    links = [
        ("Google Sheets", "https://docs.google.com/spreadsheets/d/10wtmzMIgWqPazB0yT1huNLw44W6qsM5qHhwQIYoRcqs/edit"),
        ("Азимут", "https://azimutclinic.ru"),
        ("GitHub skills", "https://github.com/lecmedea/income-agent-skills"),
        ("Telegram бот", "@Agent00AI_bot"),
    ]
    for name, url in links:
        doc.add_paragraph(f"{name}: {url}")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()